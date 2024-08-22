import os
from indox.core.document_object import Document
from typing import List

class Docx:
    """
    Load a DOCX file and extract its text and metadata, including estimated page numbers.

    Parameters:
    - file_path (str): The path to the DOCX file to be loaded.

    Methods:
    - load_file(): Extracts text from the DOCX file and returns a list of `Document` objects with associated metadata.

    Raises:
    - RuntimeError: If there is an error in loading the DOCX file.
    """

    def __init__(self, file_path: str):
        self.file_path = os.path.abspath(file_path)

    def load(self) -> List[Document]:
        import docx

        try:
            doc = docx.Document(self.file_path)
            paragraphs = doc.paragraphs

            paragraphs_per_page = 20
            num_pages = (len(paragraphs) + paragraphs_per_page - 1) // paragraphs_per_page

            # Extract text content
            documents = []
            for page in range(num_pages):
                start = page * paragraphs_per_page
                end = (page + 1) * paragraphs_per_page
                page_text = '\n'.join([p.text for p in paragraphs[start:end]])
                metadata_dict = {
                    'source': self.file_path,
                    'page': page
                }
                document = Document(metadata=metadata_dict, page_content=page_text)
                documents.append(document)

            return documents

        except Exception as e:
            raise RuntimeError(f"Error loading DOCX file: {e}")

