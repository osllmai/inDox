from typing import List, Dict, Optional
from indoxArcg.vector_stores.utils import filter_complex_metadata
from indoxArcg.core import Document


class Docx:
    """
    Load a DOCX file and extract its content and metadata.

    Parameters:
    - docx_path (str, optional): The path to the DOCX file to be loaded.
    - paragraphs_per_page (int, optional): Number of paragraphs to consider as one page. Defaults to 20.

    Methods:
    - load(docx_path): Reads the DOCX file, extracts text, and creates a `Document` object with metadata.
                      If docx_path is provided, it overrides the path set during initialization.

    Returns:
    - Document: A `Document` object with the text content and associated metadata.

    Raises:
    - FileNotFoundError: If the specified file does not exist.
    - RuntimeError: For any other errors encountered during DOCX file processing.
    """

    def __init__(self, docx_path: Optional[str] = None, paragraphs_per_page: int = 20):
        self.docx_path = docx_path
        self.content: Document = None
        self.metadata = {}
        self.paragraphs_per_page = paragraphs_per_page

    def load(self, docx_path: Optional[str] = None) -> Document:
        if docx_path:
            self.docx_path = docx_path
            
        if not self.docx_path:
            raise ValueError("No DOCX file path provided. Please provide a path during initialization or when calling load().")
            
        try:
            import docx
            from docx import Document as DocxDocument

            doc = DocxDocument(self.docx_path)
            text_content = "\n".join([para.text for para in doc.paragraphs])

            core_properties = doc.core_properties
            self.metadata = {
                prop: getattr(core_properties, prop)
                for prop in dir(core_properties)
                if not prop.startswith("_")
                and not callable(getattr(core_properties, prop))
            }

            filtered_docx = filter_complex_metadata([self])[0]
            self.metadata = filtered_docx.metadata

            self.content = Document(page_content=text_content, **self.metadata)

            return self.content

        except FileNotFoundError:
            raise FileNotFoundError(
                f"The specified file '{self.docx_path}' does not exist."
            )
        except Exception as e:
            raise RuntimeError(f"An error occurred while processing the DOCX file: {e}")

    def load_and_split(self, splitter, remove_stopwords=False):
        from indoxArcg.data_loaders.utils import load_and_process_input

        return load_and_process_input(
            loader=self.load, splitter=splitter, remove_stopwords=remove_stopwords
        )
