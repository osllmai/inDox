from typing import List, Dict
from indox.vector_stores.utils import filter_complex_metadata
from indox.core import Document

class Docx:
    """
    Load a DOCX file and extract its content and metadata.

    Parameters:
    - docx_path (str): The path to the DOCX file to be loaded.

    Methods:
    - load(): Reads the DOCX file, extracts text, and creates a `Document` object with metadata.

    Returns:
    - Document: A `Document` object with the text content and associated metadata.

    Raises:
    - FileNotFoundError: If the specified file does not exist.
    - RuntimeError: For any other errors encountered during DOCX file processing.
    """

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.content: Document = None
        self.metadata = {}

    def load(self) -> Document:
        try:
            import docx
            from docx import Document as DocxDocument

            doc = DocxDocument(self.docx_path)
            text_content = "\n".join([para.text for para in doc.paragraphs])

            core_properties = doc.core_properties
            self.metadata = {prop: getattr(core_properties, prop)
                             for prop in dir(core_properties)
                             if not prop.startswith('_') and not callable(getattr(core_properties, prop))}

            filtered_docx = filter_complex_metadata([self])[0]
            self.metadata = filtered_docx.metadata

            self.content = Document(page_content=text_content, **self.metadata)

            return self.content


        except FileNotFoundError:
            raise FileNotFoundError(f"The specified file '{self.docx_path}' does not exist.")
        except Exception as e:
            raise RuntimeError(f"An error occurred while processing the DOCX file: {e}")

    def load_and_split(self, splitter, remove_stopwords=False):
        from indox.data_loaders.utils import load_and_process_input
        return load_and_process_input(loader=self.load, splitter=splitter, remove_stopwords=remove_stopwords)

